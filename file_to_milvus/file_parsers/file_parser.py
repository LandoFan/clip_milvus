"""
文件解析器：从Word和Markdown文件中提取文本和图像
"""
import os
import io
import re
from pathlib import Path
from typing import List, Dict, Tuple
from dataclasses import dataclass

try:
    from docx import Document
    from docx.document import Document as DocumentType
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    from markdown.extensions import codehilite, tables
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

from bs4 import BeautifulSoup
from PIL import Image


@dataclass
class ExtractedContent:
    """提取的内容数据类"""
    text_chunks: List[str]  # 文本块列表
    images: List[Dict]      # 图像列表，每个包含 {'path': str, 'binary': bytes, 'format': str}
    metadata: Dict          # 元数据


class WordParser:
    """Word文档解析器"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    def parse(self, file_path: str) -> ExtractedContent:
        """
        解析Word文档
        
        Args:
            file_path: Word文档路径
            
        Returns:
            ExtractedContent对象
        """
        doc = Document(file_path)
        
        # 提取文本 - 按自然段落分块
        text_chunks = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # 每个自然段落作为一个独立的块
                text_chunks.append(text)
        
        # 提取表格文本
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text.append(row_text)
            if table_text:
                text_chunks.append("\n".join(table_text))
        
        # 提取图像
        images = []
        doc_path = Path(file_path)
        
        # 从Word文档内部提取嵌入的图像
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_part = rel.target_part
                    image_data = image_part.blob
                    
                    # 从content_type或partname获取格式
                    content_type = image_part.content_type
                    if 'png' in content_type:
                        img_format = 'PNG'
                    elif 'jpeg' in content_type or 'jpg' in content_type:
                        img_format = 'JPEG'
                    elif 'gif' in content_type:
                        img_format = 'GIF'
                    elif 'bmp' in content_type:
                        img_format = 'BMP'
                    else:
                        img_format = 'PNG'
                    
                    # 使用partname作为标识
                    img_name = image_part.partname.split('/')[-1]
                    
                    images.append({
                        'path': f"{file_path}:{img_name}",
                        'binary': image_data,
                        'format': img_format
                    })
        except Exception as e:
            print(f"Warning: Could not extract images from Word: {e}")
        
        metadata = {
            'file_path': file_path,
            'file_type': 'word',
            'title': doc.core_properties.title or Path(file_path).stem,
            'author': doc.core_properties.author or '',
            'created': str(doc.core_properties.created) if doc.core_properties.created else '',
        }
        
        return ExtractedContent(
            text_chunks=text_chunks,
            images=images,
            metadata=metadata
        )


class MarkdownParser:
    """Markdown文档解析器"""
    
    def __init__(self):
        if not MARKDOWN_AVAILABLE:
            raise ImportError("markdown is required. Install with: pip install markdown")
    
    def parse(self, file_path: str) -> ExtractedContent:
        """
        解析Markdown文档
        
        Args:
            file_path: Markdown文档路径
            
        Returns:
            ExtractedContent对象
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 转换Markdown为HTML
        md = markdown.Markdown(extensions=['codehilite', 'tables', 'fenced_code'])
        html = md.convert(md_content)
        
        # 使用BeautifulSoup解析HTML
        soup = BeautifulSoup(html, 'html.parser')
        
        # 提取文本块
        text_chunks = []
        
        # 提取段落
        for p in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            text = p.get_text().strip()
            if text and len(text) > 10:  # 过滤太短的文本
                # 按长度分块
                if len(text) > 500:
                    # 按句子分割
                    sentences = re.split(r'[.!?。！？]\s+', text)
                    chunk = ""
                    for sentence in sentences:
                        if len(chunk) + len(sentence) > 500:
                            if chunk:
                                text_chunks.append(chunk.strip())
                            chunk = sentence
                        else:
                            chunk += " " + sentence if chunk else sentence
                    if chunk:
                        text_chunks.append(chunk.strip())
                else:
                    text_chunks.append(text)
        
        # 提取代码块
        for code in soup.find_all('code'):
            text = code.get_text().strip()
            if text and len(text) > 20:
                text_chunks.append(f"[代码块]\n{text}")
        
        # 提取图像
        images = []
        md_path = Path(file_path)
        md_dir = md_path.parent
        
        # 从HTML中提取图像标签
        for img in soup.find_all('img'):
            img_src = img.get('src', '')
            if img_src:
                # 处理相对路径
                if not os.path.isabs(img_src):
                    img_path = md_dir / img_src
                else:
                    img_path = Path(img_src)
                
                if img_path.exists() and img_path.is_file():
                    try:
                        with open(img_path, 'rb') as f:
                            image_data = f.read()
                        
                        # 验证是否是有效图像
                        img_obj = Image.open(io.BytesIO(image_data))
                        img_format = img_obj.format or 'PNG'
                        
                        images.append({
                            'path': str(img_path),
                            'binary': image_data,
                            'format': img_format,
                            'alt': img.get('alt', '')
                        })
                    except Exception as e:
                        print(f"Warning: Could not load image {img_path}: {e}")
        
        # 也尝试从原始Markdown中提取图像链接
        img_pattern = r'!\[.*?\]\((.*?)\)'
        for match in re.finditer(img_pattern, md_content):
            img_src = match.group(1).strip()
            if img_src and not img_src.startswith('http'):
                if not os.path.isabs(img_src):
                    img_path = md_dir / img_src
                else:
                    img_path = Path(img_src)
                
                if img_path.exists() and img_path.is_file():
                    # 避免重复添加
                    if not any(img['path'] == str(img_path) for img in images):
                        try:
                            with open(img_path, 'rb') as f:
                                image_data = f.read()
                            img_obj = Image.open(io.BytesIO(image_data))
                            img_format = img_obj.format or 'PNG'
                            images.append({
                                'path': str(img_path),
                                'binary': image_data,
                                'format': img_format,
                                'alt': ''
                            })
                        except Exception as e:
                            print(f"Warning: Could not load image {img_path}: {e}")
        
        metadata = {
            'file_path': file_path,
            'file_type': 'markdown',
            'title': md_path.stem,
        }
        
        return ExtractedContent(
            text_chunks=text_chunks,
            images=images,
            metadata=metadata
        )


class FileParserFactory:
    """文件解析器工厂"""
    
    @staticmethod
    def get_parser(file_path: str):
        """
        根据文件扩展名获取相应的解析器
        
        Args:
            file_path: 文件路径
            
        Returns:
            相应的解析器实例
        """
        ext = Path(file_path).suffix.lower()
        
        if ext == '.docx':
            return WordParser()
        elif ext in ['.md', '.markdown']:
            return MarkdownParser()
        else:
            msg = f"Unsupported file type: {ext}"
            print(f"❌ Error: {msg}")
            raise ValueError(msg)

