"""
文件解析器单元测试
"""
import os
import sys
import tempfile
from pathlib import Path

# 添加当前目录到路径
sys.path.insert(0, str(Path(__file__).parent))

from parser.file_parser import FileParserFactory, WordParser, ExtractedContent

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def create_test_word_with_image():
    """创建一个包含图片的测试Word文档"""
    if not DOCX_AVAILABLE:
        print("❌ python-docx 未安装，无法创建测试文档")
        return None
    
    # 创建临时目录
    temp_dir = tempfile.mkdtemp()
    doc_path = os.path.join(temp_dir, "test_doc.docx")
    
    # 创建文档
    doc = Document()
    doc.add_heading('测试文档', 0)
    doc.add_paragraph('这是第一段测试文本。')
    doc.add_paragraph('这是第二段测试文本，用于验证文本提取功能。')
    
    # 创建一个简单的测试图片
    try:
        from PIL import Image
        import io
        
        # 创建一个简单的红色图片
        img = Image.new('RGB', (100, 100), color='red')
        img_path = os.path.join(temp_dir, "test_image.png")
        img.save(img_path)
        
        # 添加图片到文档
        doc.add_paragraph('下面是一张测试图片：')
        doc.add_picture(img_path, width=Inches(1.5))
        doc.add_paragraph('图片上方的文字。')
        
        print(f"✓ 创建了测试图片: {img_path}")
    except Exception as e:
        print(f"⚠ 无法添加图片: {e}")
    
    doc.save(doc_path)
    print(f"✓ 创建了测试文档: {doc_path}")
    
    return doc_path


def test_word_parser_text():
    """测试Word文档文本提取"""
    print("\n" + "=" * 60)
    print("测试1: Word文档文本提取")
    print("=" * 60)
    
    doc_path = create_test_word_with_image()
    if not doc_path:
        print("❌ 测试跳过：无法创建测试文档")
        return False
    
    try:
        factory = FileParserFactory()
        parser = factory.get_parser(doc_path)
        result = parser.parse(doc_path)
        
        print(f"\n提取结果:")
        print(f"  文本块数量: {len(result.text_chunks)}")
        print(f"  图像数量: {len(result.images)}")
        
        if result.text_chunks:
            print(f"\n文本块内容:")
            for i, chunk in enumerate(result.text_chunks, 1):
                print(f"  [{i}] {chunk[:100]}...")
        
        # 验证
        assert len(result.text_chunks) > 0, "应该提取到文本"
        print("\n✓ 文本提取测试通过")
        return True
        
    except Exception as e:
        print(f"\n❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        # 清理临时文件
        if doc_path and os.path.exists(doc_path):
            temp_dir = os.path.dirname(doc_path)
            import shutil
            shutil.rmtree(temp_dir, ignore_errors=True)


def test_word_parser_images():
    """测试Word文档图像提取"""
    print("\n" + "=" * 60)
    print("测试2: Word文档图像提取")
    print("=" * 60)
    
    doc_path = create_test_word_with_image()
    if not doc_path:
        print("❌ 测试跳过：无法创建测试文档")
        return False
    
    try:
        factory = FileParserFactory()
        parser = factory.get_parser(doc_path)
        result = parser.parse(doc_path)
        
        print(f"\n提取结果:")
        print(f"  文本块数量: {len(result.text_chunks)}")
        print(f"  图像数量: {len(result.images)}")
        
        if result.images:
            print(f"\n图像信息:")
            for i, img in enumerate(result.images, 1):
                print(f"  [{i}] 路径: {img['path']}")
                print(f"      格式: {img['format']}")
                print(f"      大小: {len(img['binary'])} bytes")
                
                # 验证图像数据有效性
                from PIL import Image
                import io
                img_obj = Image.open(io.BytesIO(img['binary']))
                print(f"      尺寸: {img_obj.size}")
        
        # 验证
        assert len(result.images) > 0, "应该提取到图像"
        assert all('binary' in img and len(img['binary']) > 0 for img in result.images), "图像应该有二进制数据"
        
        print("\n✓ 图像提取测试通过")
        return True
        
    except AssertionError as e:
        print(f"\n❌ 断言失败: {e}")
        return False
    except Exception as e:
        print(f"\n❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        # 清理临时文件
        if doc_path and os.path.exists(doc_path):
            temp_dir = os.path.dirname(doc_path)
            import shutil
            shutil.rmtree(temp_dir, ignore_errors=True)


def test_with_existing_file(file_path: str):
    """使用现有文件测试"""
    print("\n" + "=" * 60)
    print(f"测试: 解析现有文件 {file_path}")
    print("=" * 60)
    
    if not os.path.exists(file_path):
        print(f"❌ 文件不存在: {file_path}")
        return False
    
    try:
        factory = FileParserFactory()
        parser = factory.get_parser(file_path)
        result = parser.parse(file_path)
        
        print(f"\n提取结果:")
        print(f"  文本块数量: {len(result.text_chunks)}")
        print(f"  图像数量: {len(result.images)}")
        
        if result.text_chunks:
            print(f"\n前3个文本块:")
            for i, chunk in enumerate(result.text_chunks[:3], 1):
                preview = chunk[:100] + "..." if len(chunk) > 100 else chunk
                print(f"  [{i}] {preview}")
        
        if result.images:
            print(f"\n图像列表:")
            for i, img in enumerate(result.images, 1):
                print(f"  [{i}] {img['format']} - {len(img['binary'])} bytes")
        
        print("\n✓ 解析完成")
        return True
        
    except Exception as e:
        print(f"\n❌ 解析失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def run_all_tests():
    """运行所有测试"""
    print("=" * 60)
    print("文件解析器单元测试")
    print("=" * 60)
    
    results = []
    
    # 测试1: 文本提取
    results.append(("文本提取", test_word_parser_text()))
    
    # 测试2: 图像提取
    results.append(("图像提取", test_word_parser_images()))
    
    # 汇总结果
    print("\n" + "=" * 60)
    print("测试结果汇总")
    print("=" * 60)
    
    passed = 0
    failed = 0
    for name, result in results:
        status = "✓ 通过" if result else "❌ 失败"
        print(f"  {name}: {status}")
        if result:
            passed += 1
        else:
            failed += 1
    
    print(f"\n总计: {passed} 通过, {failed} 失败")
    return failed == 0


if __name__ == "__main__":
    import argparse
    
    arg_parser = argparse.ArgumentParser(description="文件解析器单元测试")
    arg_parser.add_argument('--file', type=str, help='测试指定的文件')
    args = arg_parser.parse_args()
    
    if args.file:
        test_with_existing_file(args.file)
    else:
        success = run_all_tests()
        sys.exit(0 if success else 1)

