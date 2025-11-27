"""
层次化文件解析器：支持父子分段（Hierarchical Chunking）
"""
import os
import re
from pathlib import Path
from typing import List, Dict, Optional
from dataclasses import dataclass, field
from enum import Enum


class ChunkType(Enum):
    """分段类型"""
    DOCUMENT = "document"      # 文档级别
    SECTION = "section"        # 章节级别
    SUBSECTION = "subsection"  # 小节级别
    PARAGRAPH = "paragraph"    # 段落级别
    SENTENCE = "sentence"      # 句子级别


@dataclass
class Chunk:
    """文本块数据类，支持父子关系"""
    content: str
    chunk_type: ChunkType
    index: int
    parent_id: Optional[int] = None
    children_ids: List[int] = field(default_factory=list)
    level: int = 0  # 层级深度
    metadata: Dict = field(default_factory=dict)
    
    def __hash__(self):
        return hash((self.content[:50], self.index, self.chunk_type))
    
    def __eq__(self, other):
        if not isinstance(other, Chunk):
            return False
        return (self.content[:50] == other.content[:50] and
                self.index == other.index and
                self.chunk_type == other.chunk_type)


@dataclass
class HierarchicalContent:
    """层次化内容数据类"""
    chunks: List[Chunk]  # 所有块的扁平列表
    chunk_tree: Dict[int, Chunk]  # 块索引到块的映射
    root_chunks: List[int]  # 根块索引列表
    metadata: Dict  # 文档元数据
    
    def get_chunk(self, chunk_id: int) -> Optional[Chunk]:
        """根据ID获取块"""
        return self.chunk_tree.get(chunk_id)
    
    def get_children(self, chunk_id: int) -> List[Chunk]:
        """获取子块列表"""
        chunk = self.get_chunk(chunk_id)
        if not chunk:
            return []
        return [self.chunk_tree[cid] for cid in chunk.children_ids if cid in self.chunk_tree]
    
    def get_parent(self, chunk_id: int) -> Optional[Chunk]:
        """获取父块"""
        chunk = self.get_chunk(chunk_id)
        if not chunk or chunk.parent_id is None:
            return None
        return self.chunk_tree.get(chunk.parent_id)
    
    def get_siblings(self, chunk_id: int) -> List[Chunk]:
        """获取兄弟块"""
        chunk = self.get_chunk(chunk_id)
        if not chunk or chunk.parent_id is None:
            return []
        parent = self.get_parent(chunk_id)
        if not parent:
            return []
        siblings = [self.chunk_tree[cid] for cid in parent.children_ids 
                   if cid in self.chunk_tree and cid != chunk_id]
        return siblings
    
    def get_ancestors(self, chunk_id: int) -> List[Chunk]:
        """获取所有祖先块"""
        ancestors = []
        chunk = self.get_chunk(chunk_id)
        while chunk and chunk.parent_id is not None:
            parent = self.get_parent(chunk.index)
            if parent:
                ancestors.append(parent)
                chunk = parent
            else:
                break
        return ancestors


class HierarchicalWordParser:
    """层次化Word文档解析器"""
    
    def __init__(self, 
                 max_chunk_size: int = 500,
                 min_chunk_size: int = 50,
                 overlap_size: int = 50):
        """
        初始化层次化Word解析器
        
        Args:
            max_chunk_size: 最大块大小（字符数）
            min_chunk_size: 最小块大小
            overlap_size: 块之间的重叠大小
        """
        try:
            from docx import Document
            self.Document = Document
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        self.max_chunk_size = max_chunk_size
        self.min_chunk_size = min_chunk_size
        self.overlap_size = overlap_size
    
    def parse(self, file_path: str) -> HierarchicalContent:
        """
        解析Word文档，生成层次化结构
        
        Args:
            file_path: Word文档路径
            
        Returns:
            HierarchicalContent对象
        """
        doc = self.Document(file_path)
        chunks = []
        chunk_tree = {}
        root_chunks = []
        chunk_index = 0
        
        # 第一层：文档级别
        doc_chunk = Chunk(
            content=f"Document: {Path(file_path).stem}",
            chunk_type=ChunkType.DOCUMENT,
            index=chunk_index,
            level=0,
            metadata={'file_path': file_path}
        )
        chunks.append(doc_chunk)
        chunk_tree[chunk_index] = doc_chunk
        doc_index = chunk_index
        chunk_index += 1
        
        current_section = None
        current_subsection = None
        
        # 解析段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 判断是否是标题（章节）
            style_name = para.style.name.lower() if para.style else ""
            is_heading = "heading" in style_name or para.style.name.startswith("Heading")
            
            if is_heading:
                # 提取标题级别
                heading_level = self._extract_heading_level(para.style.name)
                
                if heading_level <= 2:  # 主要章节
                    # 创建章节块
                    section_chunk = Chunk(
                        content=text,
                        chunk_type=ChunkType.SECTION,
                        index=chunk_index,
                        parent_id=doc_index,
                        level=1,
                        metadata={'heading_level': heading_level}
                    )
                    chunks.append(section_chunk)
                    chunk_tree[chunk_index] = section_chunk
                    doc_chunk.children_ids.append(chunk_index)
                    current_section = chunk_index
                    current_subsection = None
                    chunk_index += 1
                
                elif heading_level <= 4:  # 小节
                    parent_id = current_section if current_section else doc_index
                    subsection_chunk = Chunk(
                        content=text,
                        chunk_type=ChunkType.SUBSECTION,
                        index=chunk_index,
                        parent_id=parent_id,
                        level=2 if current_section else 1,
                        metadata={'heading_level': heading_level}
                    )
                    chunks.append(subsection_chunk)
                    chunk_tree[chunk_index] = subsection_chunk
                    if parent_id in chunk_tree:
                        chunk_tree[parent_id].children_ids.append(chunk_index)
                    current_subsection = chunk_index
                    chunk_index += 1
            
            else:
                # 普通段落
                # 分段处理长段落
                paragraphs = self._split_long_text(text)
                
                for para_text in paragraphs:
                    if len(para_text.strip()) < self.min_chunk_size:
                        continue
                    
                    parent_id = current_subsection if current_subsection else (
                        current_section if current_section else doc_index
                    )
                    
                    para_chunk = Chunk(
                        content=para_text,
                        chunk_type=ChunkType.PARAGRAPH,
                        index=chunk_index,
                        parent_id=parent_id,
                        level=3 if current_subsection else (2 if current_section else 1),
                        metadata={}
                    )
                    chunks.append(para_chunk)
                    chunk_tree[chunk_index] = para_chunk
                    if parent_id in chunk_tree:
                        chunk_tree[parent_id].children_ids.append(chunk_index)
                    chunk_index += 1
        
        # 处理表格
        for table_idx, table in enumerate(doc.tables):
            table_text = self._extract_table_text(table)
            if table_text:
                parent_id = current_subsection if current_subsection else (
                    current_section if current_section else doc_index
                )
                
                table_chunk = Chunk(
                    content=f"[表格 {table_idx + 1}]\n{table_text}",
                    chunk_type=ChunkType.PARAGRAPH,
                    index=chunk_index,
                    parent_id=parent_id,
                    level=3 if current_subsection else (2 if current_section else 1),
                    metadata={'table_index': table_idx}
                )
                chunks.append(table_chunk)
                chunk_tree[chunk_index] = table_chunk
                if parent_id in chunk_tree:
                    chunk_tree[parent_id].children_ids.append(chunk_index)
                chunk_index += 1
        
        root_chunks = [doc_index]
        
        metadata = {
            'file_path': file_path,
            'file_type': 'word',
            'total_chunks': len(chunks),
            'max_level': max([c.level for c in chunks]) if chunks else 0
        }
        
        return HierarchicalContent(
            chunks=chunks,
            chunk_tree=chunk_tree,
            root_chunks=root_chunks,
            metadata=metadata
        )
    
    def _extract_heading_level(self, style_name: str) -> int:
        """从样式名提取标题级别"""
        match = re.search(r'heading\s*(\d+)', style_name.lower())
        if match:
            return int(match.group(1))
        # 默认判断
        if 'heading 1' in style_name.lower():
            return 1
        elif 'heading 2' in style_name.lower():
            return 2
        elif 'heading 3' in style_name.lower():
            return 3
        return 2  # 默认级别
    
    def _split_long_text(self, text: str) -> List[str]:
        """将长文本分割成多个块"""
        if len(text) <= self.max_chunk_size:
            return [text]
        
        chunks = []
        # 按句子分割
        sentences = re.split(r'([。！？.!?]\s*)', text)
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= self.max_chunk_size:
                current_chunk += sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                # 如果单个句子就超过最大长度，强制分割
                if len(sentence) > self.max_chunk_size:
                    # 按字符分割
                    for i in range(0, len(sentence), self.max_chunk_size - self.overlap_size):
                        chunk = sentence[i:i + self.max_chunk_size]
                        if chunk.strip():
                            chunks.append(chunk.strip())
                    current_chunk = ""
                else:
                    current_chunk = sentence
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _extract_table_text(self, table) -> str:
        """提取表格文本"""
        rows_text = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                rows_text.append(row_text)
        return "\n".join(rows_text)


class HierarchicalMarkdownParser:
    """层次化Markdown文档解析器"""
    
    def __init__(self,
                 max_chunk_size: int = 500,
                 min_chunk_size: int = 50,
                 overlap_size: int = 50):
        """
        初始化层次化Markdown解析器
        
        Args:
            max_chunk_size: 最大块大小
            min_chunk_size: 最小块大小
            overlap_size: 重叠大小
        """
        try:
            import markdown
            self.markdown = markdown
        except ImportError:
            raise ImportError("markdown is required. Install with: pip install markdown")
        
        self.max_chunk_size = max_chunk_size
        self.min_chunk_size = min_chunk_size
        self.overlap_size = overlap_size
    
    def parse(self, file_path: str) -> HierarchicalContent:
        """
        解析Markdown文档，生成层次化结构
        
        Args:
            file_path: Markdown文档路径
            
        Returns:
            HierarchicalContent对象
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        chunks = []
        chunk_tree = {}
        root_chunks = []
        chunk_index = 0
        
        # 解析Markdown结构
        lines = md_content.split('\n')
        
        # 文档级别块
        doc_chunk = Chunk(
            content=f"Document: {Path(file_path).stem}",
            chunk_type=ChunkType.DOCUMENT,
            index=chunk_index,
            level=0,
            metadata={'file_path': file_path}
        )
        chunks.append(doc_chunk)
        chunk_tree[chunk_index] = doc_chunk
        doc_index = chunk_index
        chunk_index += 1
        
        current_section = None
        current_subsection = None
        current_paragraph = []
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # 检查是否是标题
            if line.startswith('#'):
                # 先处理之前的段落
                if current_paragraph:
                    para_text = '\n'.join(current_paragraph).strip()
                    if para_text and len(para_text) >= self.min_chunk_size:
                        para_chunks = self._split_long_text(para_text)
                        for para_chunk_text in para_chunks:
                            parent_id = current_subsection if current_subsection else (
                                current_section if current_section else doc_index
                            )
                            para_chunk = Chunk(
                                content=para_chunk_text,
                                chunk_type=ChunkType.PARAGRAPH,
                                index=chunk_index,
                                parent_id=parent_id,
                                level=3 if current_subsection else (2 if current_section else 1),
                                metadata={}
                            )
                            chunks.append(para_chunk)
                            chunk_tree[chunk_index] = para_chunk
                            if parent_id in chunk_tree:
                                chunk_tree[parent_id].children_ids.append(chunk_index)
                            chunk_index += 1
                    current_paragraph = []
                
                # 处理标题
                heading_level = len(line) - len(line.lstrip('#'))
                heading_text = line.lstrip('#').strip()
                
                if heading_level <= 2:  # 主要章节
                    section_chunk = Chunk(
                        content=heading_text,
                        chunk_type=ChunkType.SECTION,
                        index=chunk_index,
                        parent_id=doc_index,
                        level=1,
                        metadata={'heading_level': heading_level}
                    )
                    chunks.append(section_chunk)
                    chunk_tree[chunk_index] = section_chunk
                    doc_chunk.children_ids.append(chunk_index)
                    current_section = chunk_index
                    current_subsection = None
                    chunk_index += 1
                
                elif heading_level <= 4:  # 小节
                    parent_id = current_section if current_section else doc_index
                    subsection_chunk = Chunk(
                        content=heading_text,
                        chunk_type=ChunkType.SUBSECTION,
                        index=chunk_index,
                        parent_id=parent_id,
                        level=2 if current_section else 1,
                        metadata={'heading_level': heading_level}
                    )
                    chunks.append(subsection_chunk)
                    chunk_tree[chunk_index] = subsection_chunk
                    if parent_id in chunk_tree:
                        chunk_tree[parent_id].children_ids.append(chunk_index)
                    current_subsection = chunk_index
                    chunk_index += 1
            
            elif line.startswith('```'):  # 代码块
                # 收集代码块
                code_lines = [line]
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    code_lines.append(lines[i])
                
                code_text = '\n'.join(code_lines)
                parent_id = current_subsection if current_subsection else (
                    current_section if current_section else doc_index
                )
                
                code_chunk = Chunk(
                    content=code_text,
                    chunk_type=ChunkType.PARAGRAPH,
                    index=chunk_index,
                    parent_id=parent_id,
                    level=3 if current_subsection else (2 if current_section else 1),
                    metadata={'is_code': True}
                )
                chunks.append(code_chunk)
                chunk_tree[chunk_index] = code_chunk
                if parent_id in chunk_tree:
                    chunk_tree[parent_id].children_ids.append(chunk_index)
                chunk_index += 1
            
            elif line.strip():
                current_paragraph.append(line)
            
            else:
                # 空行，结束当前段落
                if current_paragraph:
                    para_text = '\n'.join(current_paragraph).strip()
                    if para_text and len(para_text) >= self.min_chunk_size:
                        para_chunks = self._split_long_text(para_text)
                        for para_chunk_text in para_chunks:
                            parent_id = current_subsection if current_subsection else (
                                current_section if current_section else doc_index
                            )
                            para_chunk = Chunk(
                                content=para_chunk_text,
                                chunk_type=ChunkType.PARAGRAPH,
                                index=chunk_index,
                                parent_id=parent_id,
                                level=3 if current_subsection else (2 if current_section else 1),
                                metadata={}
                            )
                            chunks.append(para_chunk)
                            chunk_tree[chunk_index] = para_chunk
                            if parent_id in chunk_tree:
                                chunk_tree[parent_id].children_ids.append(chunk_index)
                            chunk_index += 1
                    current_paragraph = []
            
            i += 1
        
        # 处理最后的段落
        if current_paragraph:
            para_text = '\n'.join(current_paragraph).strip()
            if para_text and len(para_text) >= self.min_chunk_size:
                para_chunks = self._split_long_text(para_text)
                for para_chunk_text in para_chunks:
                    parent_id = current_subsection if current_subsection else (
                        current_section if current_section else doc_index
                    )
                    para_chunk = Chunk(
                        content=para_chunk_text,
                        chunk_type=ChunkType.PARAGRAPH,
                        index=chunk_index,
                        parent_id=parent_id,
                        level=3 if current_subsection else (2 if current_section else 1),
                        metadata={}
                    )
                    chunks.append(para_chunk)
                    chunk_tree[chunk_index] = para_chunk
                    if parent_id in chunk_tree:
                        chunk_tree[parent_id].children_ids.append(chunk_index)
                    chunk_index += 1
        
        root_chunks = [doc_index]
        
        metadata = {
            'file_path': file_path,
            'file_type': 'markdown',
            'total_chunks': len(chunks),
            'max_level': max([c.level for c in chunks]) if chunks else 0
        }
        
        return HierarchicalContent(
            chunks=chunks,
            chunk_tree=chunk_tree,
            root_chunks=root_chunks,
            metadata=metadata
        )
    
    def _split_long_text(self, text: str) -> List[str]:
        """将长文本分割成多个块"""
        if len(text) <= self.max_chunk_size:
            return [text]
        
        chunks = []
        sentences = re.split(r'([。！？.!?]\s*)', text)
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= self.max_chunk_size:
                current_chunk += sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                if len(sentence) > self.max_chunk_size:
                    for i in range(0, len(sentence), self.max_chunk_size - self.overlap_size):
                        chunk = sentence[i:i + self.max_chunk_size]
                        if chunk.strip():
                            chunks.append(chunk.strip())
                    current_chunk = ""
                else:
                    current_chunk = sentence
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks

